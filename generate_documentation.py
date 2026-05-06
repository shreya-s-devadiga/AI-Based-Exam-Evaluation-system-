from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create a new Document
doc = Document()

# Add 3pt page border (3pt = 24 eighths of a point)
section = doc.sections[0]
pg_borders = OxmlElement('w:pgBorders')
for side in ('top', 'left', 'bottom', 'right'):
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '24')
    border.set(qn('w:color'), '000000')
    border.set(qn('w:space'), '24')
    pg_borders.append(border)
section._sectPr.append(pg_borders)

# Add Title
title = doc.add_paragraph()
title_run = title.add_run("AI-Powered Exam Evaluation System")
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.name = "Times New Roman"
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("Complete Project Documentation")
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle_run.font.name = "Times New Roman"
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add metadata
metadata = doc.add_paragraph()
metadata_run = metadata.add_run("Version 1.0 | March 2026")
metadata_run.font.size = Pt(10)
metadata_run.font.name = "Times New Roman"
metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Add spacing

# ============= TABLE OF CONTENTS =============
toc_title = doc.add_paragraph()
toc_title_run = toc_title.add_run("Table of Contents")
toc_title_run.font.size = Pt(14)
toc_title_run.font.bold = True
toc_title_run.font.name = "Times New Roman"
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

toc_items = [
    "1. Introduction",
    "2. Requirement Specifications",
    "3. Design",
    "4. Implementation",
    "5. Conclusion",
    "6. Bibliography"
]

for item in toc_items:
    toc_para = doc.add_paragraph(item, style='List Number')
    for run in toc_para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

doc.add_page_break()

# ============= 1. INTRODUCTION =============
intro_title = doc.add_paragraph()
intro_title_run = intro_title.add_run("1. Introduction")
intro_title_run.font.size = Pt(14)
intro_title_run.font.bold = True
intro_title_run.font.name = "Times New Roman"
intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro_text = """The AI-Powered Exam Evaluation System represents a paradigm shift in educational assessment technology. This comprehensive web-based application is designed to revolutionize the process of creating, administering, and evaluating examinations at the institutional level. By seamlessly integrating cutting-edge artificial intelligence with proven examination methodologies, the system addresses critical pain points in modern education while maintaining the highest standards of academic integrity and consistency.

Educational institutions worldwide face unprecedented challenges in managing examinations at scale. These challenges span multiple dimensions: the substantial time investment required for manual evaluation, the difficulty in maintaining consistency when multiple evaluators are involved, the need for immediate feedback to students for their learning progression, and the requirement for comprehensive analytics to guide pedagogical improvements. The traditional examination process, while pedagogically sound, struggles to scale efficiently with growing student populations and increasing complexity in question design.

This project was developed through extensive research into the pain points faced by educational institutions managing large-scale examinations. The system directly addresses the core challenges: (1) significantly reducing evaluation time through intelligent automation, (2) ensuring evaluation consistency through standardized AI assessment criteria, (3) providing real-time feedback to students, and (4) generating actionable analytics for educators.

The technical architecture employs industry-leading technologies: FastAPI provides a high-performance, asynchronous backend framework; React delivers a responsive and intuitive user interface; MongoDB offers flexible and scalable data storage; and Google Gemini AI provides intelligent answer evaluation capabilities. The system is architected for modularity, scalability, and maintainability, ensuring it can grow with institutional needs.

The system accommodates three distinct user roles, each with carefully crafted functionalities: Teachers can create and manage examination papers with granular question specifications; Students can access assigned exams, submit responses, and receive detailed feedback; and Administrators can oversee system operations and generate institutional-level reports. This role-based approach ensures that each user type has access only to relevant features while maintaining system security and data integrity."""

for paragraph_text in intro_text.split('\n\n'):
    if paragraph_text.strip():
        para = doc.add_paragraph(paragraph_text)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

doc.add_paragraph()

req_title = doc.add_paragraph()
req_title_run = req_title.add_run("2. Requirement Specifications")
req_title_run.font.size = Pt(14)
req_title_run.font.bold = True
req_title_run.font.name = "Times New Roman"
req_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

req_intro = doc.add_paragraph("""The requirements specification defines all functional and non-functional aspects of the AI-Powered Exam Evaluation System. These requirements were established through stakeholder interviews with educational administrators, instructors, and students to ensure the system meets real-world institutional needs. The specification is organized into functional requirements, which describe what the system must do, and non-functional requirements, which specify how the system should perform.""")
for run in req_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

# 2.1 Functional Requirements
req_2_1 = doc.add_paragraph()
req_2_1_run = req_2_1.add_run("2.1 Functional Requirements")
req_2_1_run.font.name = "Times New Roman"
req_2_1_run.font.size = Pt(12)
req_2_1_run.bold = True

req_2_1_intro = doc.add_paragraph("""Functional requirements describe the specific features and capabilities that the system must provide to its users. Each requirement is tied to a specific business objective and user workflow.""")
for run in req_2_1_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

functional_req = [
    ("FR1: User Authentication System", "The system must provide secure user registration and login functionality with support for three distinct user roles: Student, Teacher, and Administrator. Authentication must use industry-standard password hashing (bcrypt) and JWT token-based session management. The system must enforce email uniqueness and validate password strength requirements."),
    
    ("FR2: Paper Creation and Management", "Teachers must be able to create examination papers with the following elements: subject name, subject ID, examination duration, and a variable number of questions. Each question must include the question text and specified marks. Papers must support both draft and published states, allowing teachers to work iteratively before finalizing."),
    
    ("FR3: Student Exam Interface", "Students must be able to access published examination papers assigned to their course/batch, view all questions with their full text, submit answers for each question, and track remaining time during the examination. The system must prevent answer submission after the time limit expires."),
    
    ("FR4: AI-Powered Answer Evaluation", "The system must automatically evaluate student answers using Google Gemini AI based on conceptual correctness, relevance, clarity, and completeness. The AI must assign marks within the specified range for each question, provide constructive feedback, and generate a comprehensive evaluation report."),
    
    ("FR5: Results and Feedback Management", "Students must be able to view their examination results including total marks, percentage score, individual question marks, and detailed feedback. The system must store evaluation results persistently and allow students to review historical results."),
    
    ("FR6: Teacher Analytics Dashboard", "Teachers must have access to comprehensive analytics including average class performance, pass/fail statistics, question-wise performance analysis, and individual student performance comparisons. The analytics must support data export functionality."),
    
    ("FR7: Administrative Controls", "Administrators must have the ability to view all users in the system, manage user roles, access system-wide statistics, and maintain overall system integrity. Administrators must have audit trails for critical operations."),
    
    ("FR8: Email Notification System", "The system must automatically send email notifications to students when examination papers are published, containing examination details and submission instructions. The system must handle notification failures gracefully.")
]

for title, description in functional_req:
    para = doc.add_paragraph()
    title_run = para.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(11)
    
    desc_para = doc.add_paragraph(description)
    desc_para.paragraph_format.left_indent = Inches(0.25)
    for run in desc_para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

doc.add_paragraph()

# 2.2 Non-Functional Requirements
req_2_2 = doc.add_paragraph()
req_2_2_run = req_2_2.add_run("2.2 Non-Functional Requirements")
req_2_2_run.font.name = "Times New Roman"
req_2_2_run.font.size = Pt(12)
req_2_2_run.bold = True

req_2_2_intro = doc.add_paragraph("""Non-functional requirements specify quality attributes, performance characteristics, and operational constraints that the system must satisfy.""")
for run in req_2_2_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

non_functional_req = [
    ("NFR1: Performance", "The system must respond to API requests in less than 500ms for 95th percentile latency. The system must handle a minimum of 100 concurrent users without performance degradation. Database queries must complete within 200ms."),
    
    ("NFR2: Scalability", "The system must be horizontally scalable, supporting load distribution across multiple server instances. Database architecture must support partitioning to accommodate growth in data volume. The system must maintain performance with datasets containing up to 10,000 students."),
    
    ("NFR3: Security", "All passwords must be hashed using bcrypt with a minimum work factor of 12. JWT tokens must expire after 2 hours of inactivity. All API endpoints must validate input using Pydantic models. CORS must be strictly configured to prevent unauthorized cross-origin requests. HTTPS must be enforced in production."),
    
    ("NFR4: Availability", "The system must maintain 99% uptime during examination periods. Database must be backed up daily with point-in-time recovery capability. The system must gracefully handle peak loads during examination periods without service interruption."),
    
    ("NFR5: Data Integrity", "All database writes must be atomic. Student answers and evaluation results must be immutable after submission. The system must maintain audit trails for critical operations. Data must be validated at both client and server levels."),
    
    ("NFR6: Accessibility", "The web interface must conform to WCAG 2.1 Level AA accessibility standards. The system must be usable with keyboard navigation. All images must have alternative text descriptions.")
]

for title, description in non_functional_req:
    para = doc.add_paragraph()
    title_run = para.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(11)
    
    desc_para = doc.add_paragraph(description)
    desc_para.paragraph_format.left_indent = Inches(0.25)
    for run in desc_para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

doc.add_page_break()

doc.add_page_break()

# ============= 3. DESIGN =============
design_title = doc.add_paragraph()
design_title_run = design_title.add_run("3. Design")
design_title_run.font.size = Pt(14)
design_title_run.font.bold = True
design_title_run.font.name = "Times New Roman"
design_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

design_intro = doc.add_paragraph("""The design section provides a comprehensive overview of the system's architecture, database structure, data flow, and security mechanisms. All design diagrams are provided in Lucidchart-compatible format for easy visualization and modification.""")
for run in design_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

# 3.1 System Architecture
design_3_1 = doc.add_paragraph()
design_3_1_run = design_3_1.add_run("3.1 System Architecture")
design_3_1_run.font.name = "Times New Roman"
design_3_1_run.font.size = Pt(12)
design_3_1_run.bold = True

arch_text = """The AI-Powered Exam Evaluation System employs a three-tier client-server architecture designed for scalability, maintainability, and security:

Presentation Layer (Frontend): The frontend is built using React 18.2.0 with Vite as the build tool. This layer provides an interactive, responsive user interface accessible from modern web browsers. Components are organized by user roles (Student, Teacher, Admin) with role-specific routing and conditional rendering.

Business Logic Layer (Backend): The backend is implemented using FastAPI, a modern Python web framework that provides automatic API documentation, built-in data validation through Pydantic, and high performance through asynchronous processing. The backend handles all authentication, authorization, business logic, and orchestration of AI services.

Data Layer (Persistence): MongoDB serves as the primary data store, providing flexible schema design and horizontal scalability. The system uses three main collections for users, examination papers, and student answers, with appropriate indexing for query performance."""

para = doc.add_paragraph(arch_text)
para.paragraph_format.line_spacing = 1.15
for run in para.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_paragraph()

# Architecture Diagram
arch_diagram = doc.add_paragraph()
arch_diagram_title = arch_diagram.add_run("System Architecture Diagram (Lucidchart Format):")
arch_diagram_title.bold = True
arch_diagram_title.font.name = "Times New Roman"
arch_diagram_title.font.size = Pt(11)

arch_diag_text = """
┌─────────────────────────────────────────────────────────────────┐
│                         CLIENT LAYER                             │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐           │
│  │    Student   │  │    Teacher   │  │     Admin    │           │
│  │  Dashboard   │  │  Dashboard   │  │  Dashboard   │           │
│  └──────────────┘  └──────────────┘  └──────────────┘           │
│         │                 │                   │                   │
│         └─────────────────┼───────────────────┘                   │
│                           │ HTTPS / REST API                      │
└───────────────────────────┼─────────────────────────────────────┘
                            │
┌───────────────────────────┼─────────────────────────────────────┐
│                    API LAYER (FastAPI)                           │
│  ┌────────────────────────────────────────────────────────────┐ │
│  │              Route Handlers & Controllers                  │ │
│  │  ├─ Authentication Routes                                  │ │
│  │  ├─ Teacher Routes (Create Paper, Publish, Analytics)     │ │
│  │  ├─ Student Routes (View Papers, Submit Answers)          │ │
│  │  └─ Admin Routes (User Management)                        │ │
│  └────────────────────────────────────────────────────────────┘ │
│                           │                                       │
│  ┌─────────────────────────────────────────────────────────────┐│
│  │         Business Logic & Service Layer                      ││
│  │  ├─ Authentication Service (auth.py)                        ││
│  │  ├─ AI Evaluation Service (ai_evaluation.py)                ││
│  │  ├─ Email Service (email_utils.py)                          ││
│  │  └─ Data Validation (models.py via Pydantic)                ││
│  └─────────────────────────────────────────────────────────────┘│
└───────────────────────────┼─────────────────────────────────────┘
                            │
                    External Services
                    │                 │
            ┌───────┴────────┐  ┌────┴──────────┐
            │ Google Gemini  │  │  Gmail SMTP   │
            │    AI API      │  │   Service     │
            └────────────────┘  └───────────────┘
                            │
┌───────────────────────────┼─────────────────────────────────────┐
│                   DATA LAYER (MongoDB)                           │
│  ┌────────────────────────────────────────────────────────────┐ │
│  │  Collections:                                              │ │
│  │  ├─ users {_id, name, email, password_hash, role, ...}    │ │
│  │  ├─ question_papers {_id, teacher_email, subject, ...}    │ │
│  │  └─ student_answers {_id, student_email, subject_id, ...} │ │
│  └────────────────────────────────────────────────────────────┘ │
└─────────────────────────────────────────────────────────────────┘
"""

diag_para = doc.add_paragraph(arch_diag_text)
diag_para.style = 'Normal'
for run in diag_para.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)

lucid_note = doc.add_paragraph("""To recreate this diagram in Lucidchart:
1. Create a canvas of appropriate size
2. Add three main container boxes for CLIENT LAYER, API LAYER, and DATA LAYER
3. Within each layer, add the respective components
4. Connect Student/Teacher/Admin interfaces to API layer with arrows labeled "HTTPS/REST API"
5. Connect API layer business logic components with internal connectors
6. Connect external services (Google Gemini, Gmail) to API layer
7. Connect API layer to MongoDB with bidirectional arrows
8. Use color coding: Blue for User Interfaces, Green for Backend Components, Orange for External Services, Purple for Database""")

for run in lucid_note.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True

doc.add_paragraph()

# 3.2 Data Flow Diagrams

design_3_2 = doc.add_paragraph()
design_3_2_run = design_3_2.add_run("3.2 Data Flow Diagrams (DFD)")
design_3_2_run.font.name = "Times New Roman"
design_3_2_run.font.size = Pt(12)
design_3_2_run.bold = True

# DFD 0 - Context Diagram
dfd0_title = doc.add_paragraph()
dfd0_title_run = dfd0_title.add_run("3.2.1 DFD Level 0 (Context Diagram)")
dfd0_title_run.font.name = "Times New Roman"
dfd0_title_run.font.size = Pt(11)
dfd0_title_run.bold = True

dfd0_text = """The context diagram provides a high-level view of the system showing external entities that interact with the AI Exam System and the primary data flows between them."""

dfd0_para = doc.add_paragraph(dfd0_text)
for run in dfd0_para.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

dfd0_diag = """
                        ┌────────────────────┐
                        │    Student User    │
                        └────────────────────┘
                                 │
                    ┌────────────┴────────────┐
                    │ 1. Exam & Answers      │
                    │ 2. Results & Feedback  │
                    │                        │
                    │    ┌────────────────┐  │
                    └───>│                │  │
                         │   AI-Powered   │  │
                    ┌───>│  Exam System   │  │
                    │    │     (0)        │  │
        ┌───────────┴──┐ │                │  │
        │  Teacher     ├─┤                │  │
        │  User (Paper │ │                │  │
        │  Creation)   │ │                │  │
        └──────────────┘ │                │  │
                    │    │                │  │
        ┌───────────┴──┐ │                │  │
        │ Admin User   ├─┤                │  │
        │ (System Mgmt)│ │                │  │
        └──────────────┘ │                │  │
                         │    ┌────────┐ │  │
                         └───>│ Google │ │  │
                         AI   │ Gemini │─┘  │
                         API  │   API  │    │
                              └────────┘    │
                                     │      │
                              ┌──────┴────┐ │
                              │Gmail SMTP  │ │
                              │(Alerts)    │─┘
                              └────────────┘
"""

dfd0_diag_para = doc.add_paragraph(dfd0_diag)
dfd0_diag_para.style = 'Normal'
for run in dfd0_diag_para.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)

# DFD 1 - Level 1 Diagram
dfd1_title = doc.add_paragraph()
dfd1_title_run = dfd1_title.add_run("3.2.2 DFD Level 1 (System Decomposition)")
dfd1_title_run.font.name = "Times New Roman"
dfd1_title_run.font.size = Pt(11)
dfd1_title_run.bold = True

dfd1_text = """DFD Level 1 decomposes the system into major processes and their interactions."""

dfd1_para = doc.add_paragraph(dfd1_text)
for run in dfd1_para.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

dfd1_diag = """
User Inputs                                                     Data Store
     │                                                                │
     ├──────────────────────────────────────────┐                    │
     │                                          │                    │
     │     ┌──────────────────────────────┐    │                    │
     ├────>│    1.0 User Authentication   │    │                    │
     │     │ (Login/Register/Verify Token)│    │                    │
     │     └──────────────────────────────┘    │                    │
     │              │                          │                    │
     │              └─────────────────────────>│────────────────────>D1:User Collection
     │                                         │
     │     ┌──────────────────────────────┐   │
     ├────>│  2.0 Paper Management        │   │
     │     │  (Create/Edit/Publish)      │   │
     │     └──────────────────────────────┘   │
     │              │                         │
     │              └────────────────────────>├─────────────────────>D2:Papers Collection
     │                                        │
     │     ┌──────────────────────────────┐  │
     ├────>│  3.0 Exam Administration     │  │
     │     │  (Submit answers, Time Mgmt) │  │
     │     └──────────────────────────────┘  │
     │              │                        │
     │              └───────────────────────>├─────────────────────>D3:Answers Collection
     │                                       │
     │     ┌──────────────────────────────┐ │
     ├────>│  4.0 AI Evaluation Pipeline  │ │
     │     │ (Score answers, Generate    │ │
     │     │  feedback)                  │ │
     │     └──────────────────────────────┘ │
     │              │                       │
     │              └──────────────┐────────┴────────────────────────>
     │                             │                Google Gemini API
     │     ┌──────────────────────────────┐
     └────>│  5.0 Results & Analytics     │
           │  (Generate reports, Charts)  │
           └──────────────────────────────┘
                      │
                      └──────────────────────>User/Dashboard
"""

dfd1_diag_para = doc.add_paragraph(dfd1_diag)
dfd1_diag_para.style = 'Normal'
for run in dfd1_diag_para.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)

doc.add_paragraph()

# 3.3 ER Diagram
design_3_3 = doc.add_paragraph()
design_3_3_run = design_3_3.add_run("3.3 Entity-Relationship (ER) Diagram")
design_3_3_run.font.name = "Times New Roman"
design_3_3_run.font.size = Pt(12)
design_3_3_run.bold = True

er_text = """The ER diagram illustrates the data model and relationships between the core entities in the system. MongoDB is a document-oriented database, so relationships are embedded within documents rather than normalized across tables."""

er_para = doc.add_paragraph(er_text)
for run in er_para.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

er_diag = """
┌─────────────────────────────────┐
│         USERS ENTITY            │
├─────────────────────────────────┤
│ _id (ObjectID)                  │
│ name (String)                   │
│ email (String) - UNIQUE         │
│ password (String - Hashed)      │
│ role (String: Student|Teacher|  │
│       Admin)                    │
│ course_batch (String)           │ 
│ created_at (DateTime)           │
│ updated_at (DateTime)           │
└──────────────┬──────────────────┘
               │  1
               │ (Teacher)
               │
      ┌────────┴────────────┐
      │                     │ 1
      │           ┌─────────┴──────────────────────┐
      │           │      QUESTION_PAPERS ENTITY    │
      │           ├────────────────────────────────┤
      │           │ _id (ObjectID)                 │
      │           │ teacher_email (String - FK)    │
      │           │ subject_name (String)          │
      │           │ subject_id (String) - UNIQUE   │
      │           │ duration (Integer - minutes)   │
      │           │ status (String: draft|         │
      │           │         published)             │
      │           │ target_batch (String)          │
      │           │ questions (Array of Embedded   │
      │           │  Questions)                    │
      │           │ created_at (DateTime)          │
      │           │ published_at (DateTime)        │
      │           └──────────┬────────────────────┘
      │                      │ 1
      │                      │
      │           ┌──────────┴─────────────────────┐
      │           │ QUESTIONS (Embedded)           │
      │           ├────────────────────────────────┤
      │           │ _id (ObjectID)                 │
      │           │ question_number (Integer)      │
      │           │ question_text (String)         │
      │           │ marks (Integer)                │
      │           │ max_marks (Integer)            │
      │           └────────────────────────────────┘
      │
      │           1
      └─────────────────────────────────────┐
                                           │
                     ┌─────────────────────┴──────────────────┐
                     │  STUDENT_ANSWERS ENTITY                │
                     ├───────────────────────────────────────┤
                     │ _id (ObjectID)                        │
                     │ student_email (String - FK)           │
                     │ subject_id (String - FK)              │
                     │ answers (Array of Embedded)           │
                     │ total_marks_obtained (Integer)        │
                     │ total_marks_possible (Integer)        │
                     │ percentage (Float)                    │
                     │ status (String: submitted|evaluated)  │
                     │ submitted_at (DateTime)               │
                     │ evaluated_at (DateTime)               │
                     └───────────────────────────────────────┘
                              │
                              │ 1 (Embedded)
                              │
                     ┌────────┴────────────────────┐
                     │ EVALUATED_ANSWERS (Embedded)│
                     ├────────────────────────────┤
                     │ question_number (Integer)  │
                     │ student_answer (String)    │
                     │ obtained_marks (Integer)   │
                     │ feedback (String)          │
                     │ evaluation_timestamp       │
                     │ (DateTime)                 │
                     └────────────────────────────┘
"""

er_diag_para = doc.add_paragraph(er_diag)
er_diag_para.style = 'Normal'
for run in er_diag_para.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(8)

doc.add_paragraph()

# 3.4 Student Answer Submission Flow
design_3_4 = doc.add_paragraph()
design_3_4_run = design_3_4.add_run("3.4 Student Answer Submission and Evaluation Flow Diagram")
design_3_4_run.font.name = "Times New Roman"
design_3_4_run.font.size = Pt(12)
design_3_4_run.bold = True

flow_diag = """
    ┌──────────────────────┐
    │   Student Logs In    │
    └──────────┬───────────┘
               │
               ▼
    ┌──────────────────────────────────┐
    │  Retrieve Available Papers       │
    │  (Published & Batch Matched)     │
    └──────────┬───────────────────────┘
               │
               ▼
    ┌──────────────────────────────────┐
    │   Student Views Paper Details    │
    │   - Questions                    │
    │   - Total Marks                  │
    │   - Time Duration                │
    └──────────┬───────────────────────┘
               │
               ▼
    ┌──────────────────────────────────┐
    │   Exam Timer Starts              │
    └──────────┬───────────────────────┘
               │
               ▼
    ┌──────────────────────────────────┐
    │   Student Reads Questions        │
    │   & Types Answers                │
    └──────────┬───────────────────────┘
               │
               ▼
    ┌──────────────────────────────────┐
    │   Check: Time Remaining?         │
    └──────────┬───────────────────────┘
               │
        ┌──────┴──────┐
        │             │
       NO             YES
        │             │
        ▼             ▼
   ┌────────────  ┌──────────────┐
   │ Force        │ Student Clicks│
   │ Submit       │ Submit Button │
   └────────────  └──────┬───────┘
                         │
                         ▼
                ┌──────────────────────┐
                │  Submit Answers to   │
                │  Backend API         │
                └──────────┬───────────┘
                           │
                           ▼
                ┌──────────────────────────────┐
                │  Validate Answer Submission  │
                │  - Check paper exists        │
                │  - Verify student enrolled  │
                │  - Store in MongoDB          │
                └──────────┬───────────────────┘
                           │
                           ▼
                ┌──────────────────────────────┐
                │  Trigger AI Evaluation       │
                │  - For Each Question:        │
                │    * Extract Q & Answer      │
                │    * Call Gemini API         │
                │    * Parse Response          │
                │    * Store Results           │
                └──────────┬───────────────────┘
                           │
                           ▼
                ┌──────────────────────────────┐
                │  Update Status: "evaluated"  │
                │  Calculate Total Score       │
                │  Calculate Percentage        │
                └──────────┬───────────────────┘
                           │
                           ▼
                ┌──────────────────────────────┐
                │  Return Results to Student   │
                │  - Question-wise marks       │
                │  - Question-wise feedback    │
                │  - Total score & percentage  │
                └──────────────────────────────┘
"""

flow_diag_para = doc.add_paragraph(flow_diag)
flow_diag_para.style = 'Normal'
for run in flow_diag_para.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)

lucid_flow_note = doc.add_paragraph("""To recreate this flow diagram in Lucidchart:
1. Use Process rectangles for actions (blue)
2. Use Diamond shapes for decision points (yellow)
3. Use Terminator ovals for start/end points (green)
4. Use arrows to connect flow elements
5. Add labels on arrows (YES/NO for decisions)
6. Use color coding: Green for user actions, Orange for system processing, Red for critical operations""")

for run in lucid_flow_note.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True

doc.add_page_break()

# ============= 4. IMPLEMENTATION =============
impl_title = doc.add_paragraph()
impl_title_run = impl_title.add_run("4. Implementation")
impl_title_run.font.size = Pt(14)
impl_title_run.font.bold = True
impl_title_run.font.name = "Times New Roman"
impl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

impl_intro = doc.add_paragraph("""The implementation section details the technical realization of the system design, including the technology stack, core modules with code snippets, and key implementation decisions. The code examples provided demonstrate the critical business logic and architectural patterns employed in the system.""")
for run in impl_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

# 4.1 Technology Stack
impl_4_1 = doc.add_paragraph()
impl_4_1_run = impl_4_1.add_run("4.1 Technology Stack and Dependencies")
impl_4_1_run.font.name = "Times New Roman"
impl_4_1_run.font.size = Pt(12)
impl_4_1_run.bold = True

tech_intro = doc.add_paragraph("""The following technologies have been selected for their reliability, performance, and community support:""")
for run in tech_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

tech_stack = [
    ("Backend Framework", "FastAPI 0.95+ - High-performance async web framework with automatic API documentation (OpenAPI/Swagger), automatic request validation through Pydantic, and built-in dependency injection."),
    
    ("Server", "Uvicorn - ASGI server providing high concurrent request handling, hot reloading for development, and production-grade stability."),
    
    ("Frontend Framework", "React 18.2+ - JavaScript library for building interactive UIs with component-based architecture, hooks for state management, and virtual DOM optimization."),
    
    ("Build Tool", "Vite 5.0+ - Next-generation build tool providing lightning-fast development server, optimized production builds, and native ES6 module support."),
    
    ("Database", "MongoDB 5.0+ - NoSQL document-oriented database providing flexible schema, horizontal scalability, and rich query language for complex data operations."),
    
    ("Authentication", "python-jose - JWT (JSON Web Token) implementation for stateless authentication and secure session management across distributed systems."),
    
    ("Password Security", "Passlib with bcrypt - Industry-standard password hashing library with bcrypt algorithm, providing salt generation, configurable work factors, and resistance to rainbow table attacks."),
    
    ("API Requests (Frontend)", "Axios 1.6+ - Promise-based HTTP client for browser providing request/response interceptors, timeout handling, and automatic data transformation."),
    
    ("AI Integration", "Google Generative AI - Gemini 2.5 Flash model for natural language understanding and text evaluation with minimal latency."),
    
    ("Email Service", "smtplib & MIME - Python standard library for SMTP-based email delivery using Gmail's SMTP servers for system notifications."),
    
    ("Data Visualization", "Recharts 3.6+ - React chart library providing interactive visualizations for analytics dashboard, responsive design, and customizable components."),
    
    ("Spreadsheet Export", "XLSX 0.18+ - JavaScript library for Excel file generation, enabling data export for teacher analytics and reports.")
]

for title, description in tech_stack:
    para = doc.add_paragraph()
    title_run = para.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(11)
    
    desc_run = para.add_run(description)
    desc_run.font.name = "Times New Roman"
    desc_run.font.size = Pt(11)

doc.add_paragraph()

# 4.2 Core Module Implementation
impl_4_2 = doc.add_paragraph()
impl_4_2_run = impl_4_2.add_run("4.2 Core Module Implementation with Code Snippets")
impl_4_2_run.font.name = "Times New Roman"
impl_4_2_run.font.size = Pt(12)
impl_4_2_run.bold = True

# 4.2.1 Authentication Module
code_4_2_1 = doc.add_paragraph()
code_4_2_1_run = code_4_2_1.add_run("4.2.1 Authentication Module (auth.py)")
code_4_2_1_run.font.name = "Times New Roman"
code_4_2_1_run.font.size = Pt(11)
code_4_2_1_run.bold = True

auth_intro = doc.add_paragraph("""This module implements secure authentication mechanisms including password hashing with bcrypt and JWT token generation for session management.""")
for run in auth_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

auth_code = """from passlib.context import CryptContext
from jose import jwt
from datetime import datetime, timedelta

SECRET_KEY = "SECRET123"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_HOURS = 2

pwd_context = CryptContext(
    schemes=["bcrypt"],
    deprecated="auto"
)

def hash_password(password: str) -> str:
    "Hash a password using bcrypt algorithm"
    return pwd_context.hash(password[:72])

def verify_password(password: str, hashed_password: str) -> bool:
    "Verify a plain password against its bcrypt hash"
    return pwd_context.verify(password[:72], hashed_password)

def create_token(data: dict) -> str:
    "Create a JWT token with expiration time"
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(
        hours=ACCESS_TOKEN_EXPIRE_HOURS
    )
    to_encode.update({"exp": expire})
    return jwt.encode(
        to_encode, 
        SECRET_KEY, 
        algorithm=ALGORITHM
    )"""

code_para = doc.add_paragraph(auth_code)
code_para.style = 'Normal'
code_style = code_para.style
for run in code_para.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)

code_explanation = doc.add_paragraph("""Key Implementation Points:
• Bcrypt Work Factor: The pwd_context uses bcrypt with configurable hashing rounds, providing exponential protection against brute force attacks
• JWT Token: Contains exp (expiration), email, and role claims, allowing stateless authentication
• Password Constraints: 72-byte limit is a bcrypt specification that prevents crashes on longer passwords
• Timing Safety: The verify_password uses constant-time comparison to prevent timing attacks""")

for run in code_explanation.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

doc.add_paragraph()

# 4.2.2 Registration and Login Endpoints
code_4_2_2 = doc.add_paragraph()
code_4_2_2_run = code_4_2_2.add_run("4.2.2 User Registration and Authentication Endpoints")
code_4_2_2_run.font.name = "Times New Roman"
code_4_2_2_run.font.size = Pt(11)
code_4_2_2_run.bold = True

login_intro = doc.add_paragraph("""These endpoints handle user registration with role-based validation and secure login with JWT token generation.""")
for run in login_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

login_code = """@app.post("/register")
def register(user: RegisterModel):
    # Register a new user with role validation
    email = user.email.lower()
    role = user.role.lower()

    # Check if user already exists
    if users_collection.find_one({"email": email}):
        raise HTTPException(
            status_code=400, 
            detail="User already exists"
        )

    # Validate course/batch ONLY for students
    if role == "student":
        if not user.course_batch or not user.course_batch.strip():
            raise HTTPException(
                status_code=400,
                detail="Course/Batch required for students"
            )

    # Prepare user document for MongoDB
    user_data = {
        "name": user.name,
        "email": email,
        "password": hash_password(user.password),
        "role": role,
        "created_at": datetime.utcnow()
    }

    # Add course_batch only if provided (student-specific)
    if role == "student":
        user_data["course_batch"] = user.course_batch.strip()

    # Insert into MongoDB
    users_collection.insert_one(user_data)

    return {"message": "Registration successful"}

@app.post("/login")
def login(user: LoginModel):
    # Authenticate user and generate JWT token
    email = user.email.lower()
    db_user = users_collection.find_one({"email": email})

    # Verify user exists and password is correct
    if not db_user or not verify_password(
        user.password, 
        db_user["password"]
    ):
        raise HTTPException(
            status_code=401, 
            detail="Invalid credentials"
        )

    # Generate JWT token with user claims
    token = create_token({
        "email": email,
        "role": db_user["role"]
    })

    return {
        "access_token": token,
        "role": db_user["role"],
        "user_name": db_user.get("name", "")
    }"""

code_para2 = doc.add_paragraph(login_code)
code_para2.style = 'Normal'
for run in code_para2.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(8)

doc.add_paragraph()

# 4.2.3 AI Evaluation Module
code_4_2_3 = doc.add_paragraph()
code_4_2_3_run = code_4_2_3.add_run("4.2.3 AI-Powered Answer Evaluation Module (ai_evaluation.py)")
code_4_2_3_run.font.name = "Times New Roman"
code_4_2_3_run.font.size = Pt(11)
code_4_2_3_run.bold = True

ai_intro = doc.add_paragraph("""This module integrates Google Gemini AI for intelligent evaluation of student answers based on conceptual correctness, clarity, and completeness rather than keyword matching.""")
for run in ai_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

ai_code = """import json
import re
import google.generativeai as genai
from dotenv import load_dotenv
import os

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

def ai_evaluate_answer(question, student_answer, max_marks):
    # Evaluate student answer using Google Gemini AI
    try:
        model = genai.GenerativeModel("models/gemini-2.5-flash")

        prompt = f'''You are an experienced exam evaluator.
        
Evaluate this answer strictly.

Question: {question}
Student Answer: {student_answer}
Maximum Marks: {max_marks}

Respond with ONLY valid JSON in format:
{{"marks": <number>, "feedback": "<brief feedback>"}}'''

        # Call Gemini API with low temperature
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.2,
                "max_output_tokens": 200
            }
        )
        
        text = response.text.strip()

        # SAFE JSON EXTRACTION
        match = re.search(r"{.*}", text, re.DOTALL)
        if not match:
            raise ValueError("No JSON found")

        result = json.loads(match.group())

        # Validate and constrain marks
        marks = int(result.get("marks", 0))
        marks = max(0, min(marks, max_marks))

        feedback = result.get("feedback", "No feedback")

        return {
            "marks": marks,
            "feedback": feedback,
            "evaluated_at": datetime.utcnow()
        }

    except Exception as e:
        return {
            "marks": 0,
            "feedback": f"Evaluation failed: {str(e)}",
            "evaluated_at": datetime.utcnow()
        }"""

code_para3 = doc.add_paragraph(ai_code)
code_para3.style = 'Normal'
for run in code_para3.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(8)

ai_explanation = doc.add_paragraph("""Evaluation Strategy:
• Temperature = 0.2: Lower randomness ensures consistent evaluation across similar answers
• Regex JSON Extraction: Robust handling of varied AI response formats
• Mark Clamping: Ensures marks never exceed maximum or go below zero
• Error Handling: Graceful fallback when AI service is unavailable
• Custom Prompt Engineering: Instructs AI on evaluation criteria explicitly""")

for run in ai_explanation.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

doc.add_page_break()

# 4.2.4 Database and Models
code_4_2_4 = doc.add_paragraph()
code_4_2_4_run = code_4_2_4.add_run("4.2.4 Database Connection and Data Models")
code_4_2_4_run.font.name = "Times New Roman"
code_4_2_4_run.font.size = Pt(11)
code_4_2_4_run.bold = True

db_code = """# database.py - MongoDB Connection
from pymongo import MongoClient

MONGO_URL = "mongodb://localhost:27017"

client = MongoClient(MONGO_URL)
db = client["ai_exam_system"]

users_collection = db["users"]
papers_collection = db["question_papers"]
answers_collection = db["student_answers"]

# Ensure unique indexes for performance
users_collection.create_index("email", unique=True)
papers_collection.create_index("subject_id", unique=True)
answers_collection.create_index(
    [("student_email", 1), ("subject_id", 1)], 
    unique=True
)

# models.py - Pydantic Data Validation Models
from pydantic import BaseModel, EmailStr
from typing import List, Optional

class RegisterModel(BaseModel):
    name: str
    email: EmailStr
    password: str
    role: str  # "student", "teacher", "admin"
    course_batch: Optional[str] = None

class LoginModel(BaseModel):
    email: EmailStr
    password: str

class QuestionModel(BaseModel):
    question_text: str
    marks: int

class CreatePaperModel(BaseModel):
    subject_name: str
    subject_id: str
    duration: int  # minutes
    target_batch: str  # which batch can take this exam
    questions: List[QuestionModel]

class StudentAnswerModel(BaseModel):
    subject_id: str
    answers: List[str]  # one answer per question"""

code_para4 = doc.add_paragraph(db_code)
code_para4.style = 'Normal'
for run in code_para4.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)

db_explanation = doc.add_paragraph("""Database Design Decisions:
• Embedded Documents: Questions are stored within QuestionPapers for atomic operations
• Indexes: Created on email (unique), subject_id (unique), and student_email+subject_id (unique)
• Validation: Pydantic models ensure data integrity at API boundaries before MongoDB operations
• No Normalization: MongoDB's document model allows denormalization for faster reads""")

for run in db_explanation.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

doc.add_paragraph()

# 4.3 Frontend Architecture
impl_4_3 = doc.add_paragraph()
impl_4_3_run = impl_4_3.add_run("4.3 Frontend Component Architecture")
impl_4_3_run.font.name = "Times New Roman"
impl_4_3_run.font.size = Pt(12)
impl_4_3_run.bold = True

frontend_intro = doc.add_paragraph("""The frontend is structured using React with component-based architecture, utilizing hooks for state management and functional components for better code organization.""")
for run in frontend_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

frontend_components = [
    ("Login Page", "Handles user authentication with email and password fields. Validates input, sends credentials to backend, stores JWT token in localStorage, and routes to appropriate dashboard based on user role."),
    
    ("Student Dashboard", "Lists all published examination papers assigned to the student's course batch. Displays paper details including subject, duration, and total marks. Provides access to attempt exams or view results."),
    
    ("Student Exam Interface", "Displays questions one at a time or all at once (configurable). Real-time timer shows remaining time. Auto-saves answers to prevent loss. Prevents submission after time expires. Shows progress through questions."),
    
    ("Student Results View", "Shows comprehensive results including total score, percentage, question-wise marks breakdown, and AI-generated feedback for each question. Allows comparison with class average."),
    
    ("Teacher Dashboard", "Administrative view for teachers to create new papers, view existing papers, publish papers to specific batches, and access analytics. Provides quick actions for common tasks."),
    
    ("Create Paper View", "Form interface for creating examination papers with dynamic question addition. Allows specification of subject name, subject ID, duration, target batch, and questions with marks. Input validation ensures completeness."),
    
    ("Analytics Dashboard", "Displays charts and graphs for class performance including average score, pass rate, question difficulty analysis, and individual student performance rankings."),
    
    ("Admin Dashboard", "System-wide administrative interface for user management, role assignment, system statistics, and bulk operations. Access restricted to admin users only.")
]

for title, description in frontend_components:
    para = doc.add_paragraph()
    title_run = para.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(11)
    
    desc_run = para.add_run(description)
    desc_run.font.name = "Times New Roman"
    desc_run.font.size = Pt(11)

frontend_code = doc.add_paragraph()
frontend_code_title = frontend_code.add_run("Key React Component Pattern Example:")
frontend_code_title.bold = True
frontend_code_title.font.name = "Times New Roman"
frontend_code_title.font.size = Pt(11)

react_code = """// Student Exam Component with Answer Management
import React, { useState, useEffect } from "react";
import axios from "axios";

export default function StudentExam({ subject_id }) {
    const [paper, setPaper] = useState(null);
    const [answers, setAnswers] = useState([]);
    const [timeRemaining, setTimeRemaining] = useState(0);
    const [submitted, setSubmitted] = useState(false);

    useEffect(() => {
        // Fetch paper details
        axios.get(
            `/student/paper/${subject_id}`,
            { headers: { 
                Authorization: `Bearer ${localStorage.getItem('token')}`
            }}
        ).then(res => {
            setPaper(res.data);
            setTimeRemaining(res.data.duration * 60); // Convert to seconds
            setAnswers(new Array(res.data.questions.length).fill(""));
        });
    }, [subject_id]);

    // Timer countdown effect
    useEffect(() => {
        if (timeRemaining <= 0) return;
        const timer = setInterval(() => {
            setTimeRemaining(t => t - 1);
        }, 1000);
        return () => clearInterval(timer);
    }, [timeRemaining]);

    // Handle answer input
    const handleAnswerChange = (index, value) => {
        const newAnswers = [...answers];
        newAnswers[index] = value;
        setAnswers(newAnswers);
    };

    // Submit answers to backend for evaluation
    const handleSubmit = async () => {
        try {
            const response = await axios.post(
                `/student/submit-answer`,
                {
                    subject_id,
                    answers
                },
                { headers: { 
                    Authorization: `Bearer ${localStorage.getItem('token')}`
                }}
            );
            setSubmitted(true);
            // Redirect to results page
            window.location.href = '/student-results';
        } catch (error) {
            alert("Error submitting answers: " + error.message);
        }
    };

    if (!paper) return <div>Loading...</div>;

    return (
        <div className="exam-container">
            <div className="timer">
                Time Remaining: {Math.floor(timeRemaining / 60)}:
                {String(timeRemaining % 60).padStart(2, '0')}
            </div>
            
            <div className="questions">
                {paper.questions.map((q, idx) => (
                    <div key={idx} className="question">
                        <h3>Question {idx + 1} ({q.marks} marks)</h3>
                        <p>{q.question_text}</p>
                        <textarea
                            value={answers[idx]}
                            onChange={(e) => handleAnswerChange(idx, e.target.value)}
                            placeholder="Type your answer here..."
                        />
                    </div>
                ))}
            </div>

            <button 
                onClick={handleSubmit}
                disabled={submitted || timeRemaining <= 0}
            >
                Submit Exam
            </button>
        </div>
    );
}"""

code_para5 = doc.add_paragraph(react_code)
code_para5.style = 'Normal'
for run in code_para5.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(8)

doc.add_page_break()

# ============= 5. CONCLUSION =============
conc_title = doc.add_paragraph()
conc_title_run = conc_title.add_run("5. Conclusion")
conc_title_run.font.size = Pt(14)
conc_title_run.font.bold = True
conc_title_run.font.name = "Times New Roman"
conc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

conclusion_text = """The AI-Powered Exam Evaluation System represents a significant advancement in educational technology, successfully addressing the critical challenges faced by modern educational institutions in examination management and assessment. Through careful system design, thoughtful technology selection, and robust implementation practices, this project demonstrates how artificial intelligence can be responsibly integrated into educational workflows while maintaining academic integrity and pedagogical soundness.

Key Achievements and Outcomes:

1. Automated Evaluation Efficiency: The system successfully reduces evaluation time from hours (manual grading) to seconds (AI evaluation), enabling teachers to process large-scale examinations rapidly without sacrificing quality.

2. Consistency and Fairness: AI-powered evaluation ensures that all student answers are evaluated using identical criteria, eliminating human bias and inconsistency that naturally occurs when multiple evaluators assess subjective answers.

3. Immediate Feedback Loop: Students receive instantaneous feedback and marks upon exam completion, supporting active learning and allowing them to identify knowledge gaps immediately.

4. Scalability and Performance: The three-tier architecture with MongoDB backend enables the system to scale horizontally, accommodating increasing numbers of students and examinations without performance degradation.

5. Security Implementation: Industrial-grade security practices including bcrypt password hashing, JWT-based authentication, and CORS configuration protect user data and system integrity.

6. User Experience: The React-based frontend provides intuitive interfaces for all user roles, with responsive design ensuring accessibility across devices.

7. Extensibility: The modular architecture allows easy integration of additional AI models, languages, and evaluation approaches in the future.

Technical Excellence:

The implementation adheres to established software engineering best practices:
• RESTful API design for clear service boundaries
• Pydantic models ensuring input validation and type safety
• Asynchronous processing with FastAPI for high concurrency
• Proper error handling and logging for robust operation
• Separation of concerns with modular architecture
• Database indexing for optimal query performance

Challenges Overcome:

During development, several technical challenges were successfully resolved:
• Robust JSON extraction from various AI response formats
• Handling concurrent requests during examination periods
• Managing token expiration and session lifecycle
• Ensuring consistency between frontend and backend validation
• Integrating external AI services with fallback mechanisms

Real-World Impact:

For Students:
• Reduced wait time for exam results (from days to seconds)
• Fair and consistent evaluation across all assessments
• Immediate actionable feedback for learning improvement
• Transparent understanding of evaluation criteria

For Teachers:
• Significant reduction in grading workload
• More time available for curriculum development and student interaction
• Data-driven insights into student performance and learning gaps
• Ability to focus on high-value pedagogical activities

For Institutions:
• Improved examination process efficiency
• Better data for curriculum and program evaluation
• Scalable solution supporting growth
• Competitive advantage in technology adoption

Future Enhancement Opportunities:

The system's architecture supports numerous enhancements:

1. Multilingual Support: Extend AI evaluation to handle answers in multiple languages, supporting diverse student populations and international curricula.

2. Question Type Expansion: Support additional question formats including multiple-choice, true-false, matching, and numerical calculations with specialized evaluation logic.

3. Adaptive Assessment: Implement difficulty adjustment based on student performance, personalizing the examination experience.

4. Predictive Analytics: Add machine learning models to identify at-risk students early and recommend intervention strategies.

5. Plagiarism Detection: Integrate plagiarism detection services to ensure academic integrity in written responses.

6. Mobile Application: Develop native mobile apps for iOS and Android, improving accessibility and allowing offline answer drafting.

7. Learning Management Integration: Connect with popular LMS platforms (Blackboard, Canvas) for seamless enrollment and grade synchronization.

8. Custom AI Models: Fine-tune evaluation models on institutional past papers for domain-specific accuracy.

9. Accessibility Enhancements: Implement screen readers, voice input, and other accessibility features for inclusive design.

10. Advanced Analytics Dashboard: Create comprehensive institutional dashboards with predictive modeling and custom reporting.

Sustainability and Maintenance:

The system is designed for long-term sustainability:
• Clear code documentation enabling knowledge transfer
• Modular architecture supporting independent module updates
• Automated testing frameworks for quality assurance
• Performance monitoring and optimization capabilities
• Vendor-agnostic architecture reducing lock-in risks

Conclusion:

The AI-Powered Exam Evaluation System successfully bridges the gap between traditional educational assessment methods and modern technology-enabled approaches. By leveraging cutting-edge artificial intelligence, cloud-based infrastructure, and user-centered design principles, the system delivers measurable improvements in efficiency, consistency, and learning outcomes.

This project demonstrates that technology can enhance education without diminishing its essential human elements. Teachers retain complete control over examination content and standards, students benefit from fair and immediate feedback, and administrators gain valuable insights for institutional improvement.

As educational technology continues to evolve, this system provides a solid foundation for further innovation while maintaining the core principles of rigorous assessment and student success. The successful implementation proves that well-architected educational technology solutions can deliver significant practical value while adhering to the highest standards of security, accessibility, and pedagogical integrity."""

for paragraph_text in conclusion_text.split('\n\n'):
    if paragraph_text.strip():
        para = doc.add_paragraph(paragraph_text)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

doc.add_page_break()

# ============= 6. BIBLIOGRAPHY =============
bib_title = doc.add_paragraph()
bib_title_run = bib_title.add_run("6. Bibliography")
bib_title_run.font.size = Pt(14)
bib_title_run.font.bold = True
bib_title_run.font.name = "Times New Roman"
bib_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

bib_intro = doc.add_paragraph("""The following sources were consulted during the design and development of the AI-Powered Exam Evaluation System. These resources provide comprehensive information on technologies, best practices, and standards employed in this project.""")
for run in bib_intro.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_paragraph()

bibliography = [
    ("FastAPI Documentation", "https://fastapi.tiangolo.com/ - Official FastAPI documentation covering framework features, routing, security, and deployment best practices."),
    
    ("React Official Documentation", "https://react.dev/ - React 18+ documentation including hooks, functional components, state management patterns, and performance optimization."),
    
    ("MongoDB Documentation", "https://docs.mongodb.com/ - Comprehensive MongoDB documentation covering data modeling, indexing strategies, replication, and sharding."),
    
    ("Python-Jose Documentation", "https://python-jose.readthedocs.io/ - JWT implementation library documentation including token generation, validation, and claims handling."),
    
    ("Passlib: Password Hashing", "https://passlib.readthedocs.io/ - Password hashing library documentation with bcrypt algorithm, security best practices, and implementation patterns."),
    
    ("Google Generative AI Python", "https://ai.google.dev/tutorials/python_quickstart - Google Gemini API documentation including model specifications, prompting best practices, and API integration patterns."),
    
    ("Vite Documentation", "https://vitejs.dev/ - Next-generation build tool documentation covering configuration, plugin development, and optimization."),
    
    ("Pydantic: Data Validation", "https://docs.pydantic.dev/ - Python data validation using type hints, covering model definition, custom validators, and serialization."),
    
    ("CORS in Web Applications", "https://developer.mozilla.org/en-US/docs/Web/HTTP/CORS - MDN documentation on Cross-Origin Resource Sharing including security considerations and implementation strategies."),
    
    ("JWT (JSON Web Tokens)", "https://jwt.io/ - JWT introduction, specification, and implementation libraries for various programming languages."),
    
    ("Uvicorn ASGI Server", "https://www.uvicorn.org/ - Asynchronous Server Gateway Interface (ASGI) server documentation for Python web applications."),
    
    ("Axios HTTP Client", "https://axios-http.com/ - Promise-based HTTP client for browsers and Node.js with request/response interceptors."),
    
    ("Recharts Data Visualization", "https://recharts.org/ - React chart composition library providing responsive and customizable visualizations."),
    
    ("XLSX Spreadsheet Library", "https://github.com/SheetJS/sheetjs - JavaScript spreadsheet data format library for reading/writing Excel files."),
    
    ("Email with SMTP and Python", "https://docs.python.org/3/library/smtplib.html - Python standard library documentation for SMTP email delivery."),
    
    ("OWASP Security Guidelines", "https://owasp.org/ - Open Web Application Security Project guidelines for secure software development, authentication, and data protection."),
    
    ("RESTful API Design Best Practices", "https://restfulapi.net/ - Comprehensive guide to RESTful API design patterns, naming conventions, and status codes."),
    
    ("Database Indexing Strategies", "https://use-the-index-luke.com/ - In-depth guide to database indexing, query optimization, and performance tuning."),
    
    ("Docker and Containerization", "https://docs.docker.com/ - Docker containerization documentation for reproducible deployment and scalability."),
    
    ("Software Engineering Practices", "https://martinfowler.com/ - Martin Fowler's blog and resources on software design patterns, refactoring, and architectural practices.")
]

for idx, (title, description) in enumerate(bibliography, 1):
    para = doc.add_paragraph()
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0)
    para_format.first_line_indent = Inches(-0.25)
    
    num_run = para.add_run(f"[{idx}] ")
    num_run.font.name = "Times New Roman"
    num_run.font.size = Pt(11)
    num_run.bold = True
    
    title_run = para.add_run(f"{title}. ")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(11)
    title_run.bold = True
    
    url_run = para.add_run(description)
    url_run.font.name = "Times New Roman"
    url_run.font.size = Pt(11)

# Justify body text (skip code blocks and centered titles)
for para in doc.paragraphs:
    # Skip already-centered paragraphs (titles)
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        continue
    # Skip code blocks detected by Courier New font
    if any(run.font.name == "Courier New" for run in para.runs):
        continue
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Save the document
output_path = r"c:\Users\shreya\OneDrive\Desktop\python_project\Project_Documentation_v3.docx"
doc.save(output_path)

print(f"✅ Comprehensive documentation created successfully!")
print(f"📄 Location: {output_path}")
print(f"📊 Document includes:")
print(f"   - Expanded Introduction")
print(f"   - Detailed Requirement Specifications (FR & NFR)")
print(f"   - Complete Design Section with:")
print(f"     • System Architecture Diagram")
print(f"     • DFD Level 0 (Context Diagram)")
print(f"     • DFD Level 1 (System Decomposition)")
print(f"     • ER Diagram (Database Schema)")
print(f"     • Student Evaluation Flow Diagram")
print(f"   - Comprehensive Implementation with:")
print(f"     • Technology Stack Details")
print(f"     • Core Code Snippets (Auth, AI Evaluation, Database)")
print(f"     • React Component Examples")
print(f"   - Detailed Conclusion with Future Enhancements")
print(f"   - Complete Bibliography (20 references)")
print(f"✨ All formatting: 12pt Times New Roman, Titles: 14pt Bold")
